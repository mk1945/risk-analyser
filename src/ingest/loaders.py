from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

from pypdf import PdfReader
from docx import Document


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    filetype: str


def _load_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _load_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts).strip()


def _load_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace")


def load_document(filename: str, data: bytes) -> LoadedDocument:
    name = filename.lower().strip()
    if name.endswith(".pdf"):
        return LoadedDocument(text=_load_pdf(data), filetype="pdf")
    if name.endswith(".doc"):
        raise ValueError(
            "Unsupported legacy .doc format in local mode. Please save/convert to .docx (recommended) or export to .txt, then re-upload."
        )
    if name.endswith(".docx"):
        return LoadedDocument(text=_load_docx(data), filetype="docx")
    if name.endswith(".txt"):
        return LoadedDocument(text=_load_txt(data), filetype="txt")
    raise ValueError(f"Unsupported file type: {filename}")
